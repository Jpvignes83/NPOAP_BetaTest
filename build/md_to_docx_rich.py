from pathlib import Path
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"\*([^*]+)\*")


def _set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _add_cover_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("NPOAP")
    r.bold = True
    r.font.size = Pt(30)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run("Manuel Utilisateur")
    r2.bold = True
    r2.font.size = Pt(24)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r3 = p3.add_run("Version riche (Word)")
    r3.italic = True
    r3.font.size = Pt(13)

    doc.add_paragraph("")
    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r4 = p4.add_run(datetime.now().strftime("%d/%m/%Y"))
    r4.font.size = Pt(12)

    # Saut de page après couverture
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_word_toc(doc: Document):
    doc.add_heading("Sommaire", level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    doc.add_paragraph(
        "Cliquez dans le sommaire puis choisissez « Mettre à jour le champ » dans Word "
        "pour afficher les numéros de pages."
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_inline_runs(paragraph, text: str):
    # Parsing simple: code > bold > italic
    i = 0
    while i < len(text):
        m_code = INLINE_CODE_RE.search(text, i)
        m_bold = BOLD_RE.search(text, i)
        m_ital = ITALIC_RE.search(text, i)

        candidates = [m for m in [m_code, m_bold, m_ital] if m]
        if not candidates:
            paragraph.add_run(text[i:])
            break

        m = min(candidates, key=lambda x: x.start())
        if m.start() > i:
            paragraph.add_run(text[i:m.start()])

        content = m.group(1)
        if m.re is INLINE_CODE_RE:
            r = paragraph.add_run(content)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif m.re is BOLD_RE:
            r = paragraph.add_run(content)
            r.bold = True
        else:
            r = paragraph.add_run(content)
            r.italic = True

        i = m.end()


def convert_markdown_to_docx(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    doc = Document()
    _set_default_style(doc)

    # Couverture + sommaire Word (TOC champ)
    _add_cover_page(doc)
    _add_word_toc(doc)
    title = doc.add_heading("MANUEL UTILISATEUR NPOAP", level=0)
    if title.runs:
        title.runs[0].bold = True

    in_code = False
    code_buffer = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph()
        for idx, ln in enumerate(code_buffer):
            r = p.add_run(ln)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            if idx < len(code_buffer) - 1:
                r.add_break()
        code_buffer = []

    for raw in lines:
        line = raw.rstrip("\n")
        s = line.strip()

        if s.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not s:
            doc.add_paragraph("")
            continue

        # Séparateur horizontal approximé
        if re.match(r"^---+$", s):
            p = doc.add_paragraph()
            p.add_run("─" * 80)
            continue

        # Titres
        m_h = re.match(r"^(#{1,6})\s+(.+)$", s)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            heading_text = m_h.group(2).strip()
            doc.add_heading(heading_text, level=level)
            continue

        # Liste numérotée
        m_num = re.match(r"^(\d+)\.\s+(.+)$", s)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m_num.group(2))
            continue

        # Liste à puces
        m_b = re.match(r"^[-*]\s+(.+)$", s)
        if m_b:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m_b.group(1))
            continue

        # Paragraphe normal
        p = doc.add_paragraph()
        _add_inline_runs(p, s)

    flush_code()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    base = Path(__file__).resolve().parents[1]
    md = base / "docs" / "MANUEL_UTILISATEUR.md"
    out = base / "docs" / "MANUEL_UTILISATEUR_rich_couverture_sommaire.docx"
    convert_markdown_to_docx(md, out)
    print(out)

